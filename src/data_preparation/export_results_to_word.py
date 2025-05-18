from typing import List, Dict
from tqdm import tqdm
import os
import json
from ast import literal_eval
import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import docx.opc.constants
from datetime import datetime
import sys


# Function to convert numbers to Roman numerals for section numbering
def to_roman(num):
    val = [1000, 900, 500, 400, 100, 90, 50, 40, 10, 9, 5, 4, 1]
    syb = ["M", "CM", "D", "CD", "C", "XC", "L", "XL", "X", "IX", "V", "IV", "I"]
    roman_num = ""
    i = 0
    while num > 0:
        for _ in range(num // val[i]):
            roman_num += syb[i]
            num -= val[i]
        i += 1
    return roman_num


# Function to add a horizontal line
def add_horizontal_line(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    # Create the horizontal line as a border for the paragraph
    pBdr = OxmlElement("w:pBdr")

    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")  # Horizontal line
    bottom.set(qn("w:sz"), "6")  # Line thickness
    bottom.set(qn("w:space"), "1")  # Space above the line
    bottom.set(qn("w:color"), "auto")  # Default color

    pBdr.append(bottom)

    paragraph._p.get_or_add_pPr().append(pBdr)


# Function to handle missing information and replace it with "No information available"
def safe_str(value):
    if pd.isna(value) or value in ["", None]:
        return "No information available"
    return str(value)


# Helper function to format headings (titles only)
def format_paragraph(paragraph, bold=False, font_size=None, font_color=None):
    run = paragraph.runs[0]
    if bold:
        run.bold = True
    if font_size:
        run.font.size = Pt(font_size)
    if font_color:
        run.font.color.rgb = RGBColor(*font_color)


def _append_one_evidence(doc: Document, row: pd.Series, evidence_idx: int):

    extract = row["Extracted Entries"]
    page_number = row["page_number"]
    document_source = row["docs_path"].replace("data/docs/", "")
    relevance_score = f"{round(row['Framework Probabilities'] * 100)}%"

    analysis_paragraph = doc.add_paragraph()
    one_row_str = f"{evidence_idx}. {extract} ({document_source}, Page {page_number}, Relevance Score: {relevance_score}"
    analysis_paragraph.add_run(safe_str(one_row_str))


def _add_TOC(framework: Dict[str, Dict[str, List[str]]], toc_entries):

    for i, (tag_l1, tags_level_2) in enumerate(framework.items()):
        roman_num = to_roman(i + 1)
        toc_entries.append((f"{roman_num}. {tag_l1}", 1))  # Add to TOC with level 1

        for j, (tag_level_2_value, tags_level_3) in enumerate(tags_level_2.items()):

            arabic_counter = j + 1

            toc_entries.append(
                (f"{roman_num}.{arabic_counter}. {tag_l1} -> {tag_level_2_value}", 2)
            )  # Level 2 TOC

            for k, tag_level_3_value in enumerate(tags_level_3):
                alphabet = chr(97 + k)
                toc_entries.append(
                    (
                        f"{roman_num}.{arabic_counter}.{alphabet}. {tag_level_3_value}",
                        3,
                    )
                )


def _add_content_sections(
    doc, df, framework: Dict[str, Dict[str, List[str]]], cutoff_threshold: float
):
    
    tot_n_tags = len([s for tag, subtags in framework.items() for subsubtag in subtags.values() for s in subsubtag])
    
    tqdm_bar = tqdm(total=tot_n_tags, desc="Processing tags")

    # Continue with the rest of the content
    for i, (tag_l1, tags_level_2) in enumerate(framework.items()):
        roman_num = to_roman(i + 1)
        section_title = doc.add_heading(f"{roman_num}. {tag_l1}", level=1)
        format_paragraph(section_title, bold=True, font_size=24, font_color=(255, 0, 0))

        for k, (tag_level_2_value, tags_level_3) in enumerate(tags_level_2.items()):
            arabic_counter = k + 1
            subsection_title = doc.add_heading(
                f"{roman_num}.{arabic_counter}. {tag_l1} -> {tag_level_2_value}",
                level=2,
            )
            format_paragraph(
                subsection_title, bold=True, font_size=18, font_color=(255, 0, 0)
            )

            for alphabetic_counter, tag_level_3_value in enumerate(tags_level_3):
                alphabet = chr(97 + alphabetic_counter)
                subsubsection_title = doc.add_heading(
                    f"{roman_num}.{arabic_counter}.{alphabet}. {tag_level_3_value}",
                    level=3,
                )
                format_paragraph(
                    subsubsection_title, bold=True, font_size=16, font_color=(255, 0, 0)
                )

                tag = f"{tag_l1}->{tag_level_2_value}->{tag_level_3_value}"

                df_tag = df.copy()
                df_tag["Framework Probabilities"] = df_tag[
                    "Framework Probabilities"
                ].apply(lambda x: x[tag])
                df_tag = df_tag[
                    df_tag["Framework Probabilities"] > cutoff_threshold
                ].sort_values(by="Framework Probabilities", ascending=False).reset_index(drop=True)

                if not df_tag.empty:
                    # Add evidence entries with numbered references
                    for evidence_idx, row in df_tag.iterrows():
                        _append_one_evidence(doc, row, evidence_idx + 1)
                        
                tqdm_bar.update(1)

    tqdm_bar.close()
    
    
def generate_evidences_doc(
    analysis_file_path: os.PathLike,
    framework_file_path: os.PathLike,
    output_path: os.PathLike,
    doc_title: str,
    cutoff_threshold: float = 0.65,
):

    df = pd.read_csv(analysis_file_path)
    df["Framework Probabilities"] = df["Framework Probabilities"].apply(literal_eval)
    with open(framework_file_path, "r") as f:
        framework = json.load(f)

    # List to store TOC entries
    toc_entries = []

    # Create a Word document
    doc = Document()

    # Create the title of the document
    title = doc.add_heading(doc_title, 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add a centered "Table of Contents" title
    toc_title = doc.add_heading("Table of Contents", level=1)
    toc_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add a paragraph for the TOC
    toc_paragraph = doc.add_paragraph()
    # doc.add_page_break()

    _add_TOC(framework, toc_entries)
    doc.save(output_path)

    # Insert TOC entries after the "Table of Contents" title
    for entry, level in toc_entries:
        toc_line = toc_paragraph.add_run(entry + "\n")
        if level == 1:
            toc_paragraph.paragraph_format.left_indent = Pt(0)
            toc_line.bold = True
        elif level == 2:
            toc_paragraph.paragraph_format.left_indent = Pt(20)  # Indent level 2
        elif level == 3:
            toc_paragraph.paragraph_format.left_indent = Pt(40)  # Indent level 3

    # Ensure the TOC ends properly, and then add a page break after the TOC
    # doc.add_paragraph()
    doc.add_page_break()

    _add_content_sections(doc, df, framework, cutoff_threshold)

    # Save the Word document
    doc.save(output_path)
